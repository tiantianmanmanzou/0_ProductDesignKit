"""将 Markdown 文件转换为 DOCX 文档，保留标题层级和格式。"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def parse_markdown_line(line: str) -> Tuple[str, str]:
    """
    解析 Markdown 行，返回 (类型, 内容)。
    
    类型可能是: 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'list', 'text'
    """
    line = line.rstrip()
    
    # 匹配标题 (# 到 ######)
    heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
    if heading_match:
        level = len(heading_match.group(1))
        content = heading_match.group(2)
        return f'h{level}', content
    
    # 匹配列表项 (- 或 * 开头)
    list_match = re.match(r'^[\-\*]\s+(.+)$', line)
    if list_match:
        content = list_match.group(1)
        return 'list', content
    
    # 普通文本
    if line.strip():
        return 'text', line
    
    # 空行
    return 'empty', ''


def set_chinese_font(run, font_name: str = '宋体'):
    """设置中文字体（解决 python-docx 中文字体问题）。"""
    run.font.name = font_name
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph_with_style(
    doc: Document,
    content: str,
    style_type: str,
    font_name: str = '宋体',
    font_size: int = 12,
    bold: bool = False,
    color: RGBColor = RGBColor(0, 0, 0)
):
    """添加带样式的段落。"""
    
    # 根据类型选择样式
    style_map = {
        'h1': 'Heading 1',
        'h2': 'Heading 2',
        'h3': 'Heading 3',
        'h4': 'Heading 4',
        'h5': 'Heading 5',
        'h6': 'Heading 6',
        'list': 'List Bullet',
        'text': 'Normal',
    }
    
    style = style_map.get(style_type, 'Normal')
    para = doc.add_paragraph(content, style=style)
    
    # 设置字体样式
    for run in para.runs:
        run.font.size = Pt(font_size)
        run.font.color.rgb = color
        run.font.bold = bold
        set_chinese_font(run, font_name)
    
    return para


def convert_markdown_to_docx(
    markdown_path: Path,
    output_path: Path | None = None,
    font_name: str = '宋体',
    font_color: RGBColor = RGBColor(0, 0, 0)
) -> Path:
    """
    将 Markdown 文件转换为 DOCX 文档。
    
    Args:
        markdown_path: Markdown 文件路径
        output_path: 输出 DOCX 文件路径（默认与 Markdown 同名）
        font_name: 字体名称（默认宋体）
        font_color: 字体颜色（默认黑色）
    
    Returns:
        生成的 DOCX 文件路径
    """
    markdown_path = markdown_path.expanduser().resolve()
    if not markdown_path.exists():
        raise FileNotFoundError(f"Markdown 文件不存在: {markdown_path}")
    
    # 确定输出路径
    if output_path is None:
        output_path = markdown_path.with_suffix('.docx')
    else:
        output_path = output_path.expanduser().resolve()
    
    # 创建 Word 文档
    doc = Document()
    
    # 读取 Markdown 内容
    content = markdown_path.read_text(encoding='utf-8')
    lines = content.split('\n')
    
    # 逐行解析并添加到文档
    for line in lines:
        line_type, line_content = parse_markdown_line(line)
        
        if line_type == 'empty':
            # 空行，添加空段落
            doc.add_paragraph()
            continue
        
        # 根据类型设置字体大小和加粗
        if line_type == 'h1':
            font_size = 22
            bold = True
        elif line_type == 'h2':
            font_size = 18
            bold = True
        elif line_type == 'h3':
            font_size = 16
            bold = True
        elif line_type == 'h4':
            font_size = 14
            bold = True
        elif line_type in ('h5', 'h6'):
            font_size = 12
            bold = True
        else:
            font_size = 12
            bold = False
        
        # 添加段落
        add_paragraph_with_style(
            doc,
            line_content,
            line_type,
            font_name=font_name,
            font_size=font_size,
            bold=bold,
            color=font_color
        )
    
    # 保存文档
    doc.save(str(output_path))
    print(f"✅ 已生成 DOCX 文档: {output_path}")
    
    return output_path


def main():
    """命令行入口。"""
    parser = argparse.ArgumentParser(
        description="将 Markdown 文件转换为 DOCX 文档，保留标题层级和格式"
    )
    parser.add_argument(
        'markdown_path',
        type=Path,
        help='待转换的 Markdown 文件路径'
    )
    parser.add_argument(
        '-o', '--output',
        type=Path,
        default=None,
        help='输出 DOCX 文件路径（默认与 Markdown 同名）'
    )
    parser.add_argument(
        '-f', '--font',
        type=str,
        default='宋体',
        help='字体名称（默认：宋体）'
    )
    
    args = parser.parse_args()
    
    # 执行转换
    convert_markdown_to_docx(
        markdown_path=args.markdown_path,
        output_path=args.output,
        font_name=args.font
    )


if __name__ == '__main__':
    main()
