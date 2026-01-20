#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
è¡¨æ ¼è½¬ç»“æ„åŒ–DOCXæ–‡æ¡£
å°†è¡¨æ ¼å†…å®¹è½¬æ¢ä¸ºæ ‡é¢˜+æ­£æ–‡çš„DOCXæ–‡æ¡£æ ¼å¼
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


class TableToStructuredDocxConverter:
    """è¡¨æ ¼è½¬ç»“æ„åŒ–DOCXè½¬æ¢å™¨"""
    
    def __init__(self, input_docx_path: str):
        """
        åˆå§‹åŒ–è½¬æ¢å™¨
        
        Args:
            input_docx_path: è¾“å…¥DOCXæ–‡ä»¶è·¯å¾„
        """
        self.input_path = input_docx_path
        self.input_doc = Document(input_docx_path)
        self.output_doc = Document()
        
    def set_chinese_font(self, run, font_name='å¾®è½¯é›…é»‘'):
        """
        è®¾ç½®ä¸­æ–‡å­—ä½“
        
        Args:
            run: æ–‡æœ¬è¿è¡Œå¯¹è±¡
            font_name: å­—ä½“åç§°
        """
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def add_title(self, text: str):
        """
        æ·»åŠ æ–‡æ¡£æ ‡é¢˜
        
        Args:
            text: æ ‡é¢˜æ–‡æœ¬
        """
        title = self.output_doc.add_heading(text, level=1)
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        self.set_chinese_font(title_run)
        
    def add_heading_2(self, text: str):
        """
        æ·»åŠ äºŒçº§æ ‡é¢˜
        
        Args:
            text: æ ‡é¢˜æ–‡æœ¬
        """
        heading = self.output_doc.add_heading(text, level=2)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(31, 73, 125)
        self.set_chinese_font(heading_run)
        
    def add_heading_3(self, text: str):
        """
        æ·»åŠ ä¸‰çº§æ ‡é¢˜
        
        Args:
            text: æ ‡é¢˜æ–‡æœ¬
        """
        heading = self.output_doc.add_heading(text, level=3)
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(79, 129, 189)
        self.set_chinese_font(heading_run)
        
    def add_paragraph(self, text: str):
        """
        æ·»åŠ æ­£æ–‡æ®µè½
        
        Args:
            text: æ®µè½æ–‡æœ¬
        """
        para = self.output_doc.add_paragraph(text)
        para_run = para.runs[0]
        para_run.font.size = Pt(11)
        self.set_chinese_font(para_run)
        
    def process_table_row(self, row_cells: list, current_h1: str):
        """
        å¤„ç†è¡¨æ ¼è¡Œï¼Œè½¬æ¢ä¸ºæ ‡é¢˜+æ­£æ–‡æ ¼å¼
        
        Args:
            row_cells: å½“å‰è¡Œçš„å•å…ƒæ ¼åˆ—è¡¨
            current_h1: å½“å‰çš„ä¸€çº§æ ‡é¢˜
            
        Returns:
            æ–°çš„ä¸€çº§æ ‡é¢˜ï¼ˆå¦‚æœæœ‰å˜åŒ–ï¼‰
        """
        # è¡¨æ ¼ç»“æ„ï¼šç¬¬ä¸€åˆ—=ä¸€çº§æ ‡é¢˜ | ç¬¬äºŒåˆ—=äºŒçº§æ ‡é¢˜ | ç¬¬ä¸‰åˆ—=æ­£æ–‡
        if len(row_cells) >= 3:
            h1_text = row_cells[0].strip()
            h2_text = row_cells[1].strip()
            content = row_cells[2].strip()
            
            # å¦‚æœä¸€çº§æ ‡é¢˜å‘ç”Ÿå˜åŒ–ä¸”ä¸ä¸ºç©ºï¼Œæ·»åŠ æ–°çš„ä¸€çº§æ ‡é¢˜
            if h1_text and h1_text != current_h1:
                self.add_heading_2(h1_text)
                current_h1 = h1_text
            
            # äºŒçº§æ ‡é¢˜ï¼ˆæŠ€æœ¯æŒ‡æ ‡ï¼‰
            if h2_text:
                self.add_heading_3(h2_text)
            
            # æ­£æ–‡å†…å®¹
            if content:
                # å¤„ç†æ¢è¡Œç¬¦
                content = content.replace('<br>', '\n')
                self.add_paragraph(content)
                
        return current_h1
    
    def convert(self) -> Document:
        """
        æ‰§è¡Œè½¬æ¢
        
        Returns:
            è½¬æ¢åçš„Documentå¯¹è±¡
        """
        # æ·»åŠ æ–‡æ¡£ä¸»æ ‡é¢˜ï¼ˆä»ç¬¬ä¸€ä¸ªæ®µè½è·å–ï¼‰
        for para in self.input_doc.paragraphs:
            if para.text.strip():
                self.add_title(para.text.strip())
                break
        
        # å¤„ç†æ‰€æœ‰è¡¨æ ¼
        for table_idx, table in enumerate(self.input_doc.tables):
            if not table.rows:
                continue
            
            # è·Ÿè¸ªå½“å‰çš„ä¸€çº§æ ‡é¢˜
            current_h1 = ""
            
            # å¤„ç†æ•°æ®è¡Œï¼ˆè·³è¿‡è¡¨å¤´ï¼‰
            for row_idx, row in enumerate(table.rows[1:], start=1):
                row_cells = [cell.text.strip() for cell in row.cells]
                current_h1 = self.process_table_row(row_cells, current_h1)
        
        return self.output_doc
    
    def save_to_file(self, output_path: str):
        """
        ä¿å­˜è½¬æ¢ç»“æœåˆ°æ–‡ä»¶
        
        Args:
            output_path: è¾“å‡ºæ–‡ä»¶è·¯å¾„
        """
        self.convert()
        
        # ç¡®ä¿è¾“å‡ºç›®å½•å­˜åœ¨
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # ä¿å­˜æ–‡æ¡£
        self.output_doc.save(output_path)
        
        print(f"âœ… è½¬æ¢å®Œæˆï¼")
        print(f"ğŸ“„ è¾“å…¥æ–‡ä»¶: {self.input_path}")
        print(f"ğŸ“ è¾“å‡ºæ–‡ä»¶: {output_path}")
        print(f"ğŸ“Š æ–‡æ¡£ç»Ÿè®¡:")
        print(f"   - åŸæ–‡æ¡£æ®µè½æ•°: {len(self.input_doc.paragraphs)}")
        print(f"   - åŸæ–‡æ¡£è¡¨æ ¼æ•°: {len(self.input_doc.tables)}")
        print(f"   - æ–°æ–‡æ¡£æ®µè½æ•°: {len(self.output_doc.paragraphs)}")


def main():
    """ä¸»å‡½æ•°"""
    # è¾“å…¥æ–‡ä»¶è·¯å¾„
    input_file = "/Users/zhangxy/GenAI/DocPilot/docs/åˆ†ç±»åˆ†çº§é¡¹ç›®/æ•°æ®åˆ†ç±»åˆ†çº§å¹³å°æ‹›æ ‡éœ€æ±‚è¯´æ˜æ–‡ä»¶(1117).docx"
    
    # è¾“å‡ºæ–‡ä»¶è·¯å¾„
    output_file = input_file.replace('.docx', '_ç»“æ„åŒ–.docx')
    
    # æ£€æŸ¥è¾“å…¥æ–‡ä»¶æ˜¯å¦å­˜åœ¨
    if not os.path.exists(input_file):
        print(f"âŒ é”™è¯¯: æ–‡ä»¶ä¸å­˜åœ¨ - {input_file}")
        return
    
    try:
        # åˆ›å»ºè½¬æ¢å™¨å¹¶æ‰§è¡Œè½¬æ¢
        converter = TableToStructuredDocxConverter(input_file)
        converter.save_to_file(output_file)
        
    except Exception as e:
        print(f"âŒ è½¬æ¢å¤±è´¥: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
