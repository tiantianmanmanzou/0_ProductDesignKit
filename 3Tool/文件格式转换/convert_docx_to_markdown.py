#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX to Markdown Converter
将DOCX/DOC文档转换为Markdown格式，保留表格和标题层级关系

支持格式：
- .docx (Office Open XML)
- .docm (启用宏的Word文档)
- .doc (旧版Word二进制格式，需要系统支持textutil或libreoffice)
"""

import os
import argparse
import zipfile
import tempfile
import shutil
import subprocess
import platform
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def convert_doc_to_docx(doc_path: str) -> str:
    """
    将旧版 .doc 文件转换为 .docx 格式（临时文件）
    
    支持的转换工具（按优先级）：
    1. macOS: textutil (系统自带)
    2. 跨平台: LibreOffice (需安装)
    
    Args:
        doc_path: .doc 文件路径
        
    Returns:
        临时 .docx 文件路径
        
    Raises:
        RuntimeError: 如果没有可用的转换工具
    """
    # 创建临时目录
    temp_dir = tempfile.mkdtemp()
    temp_docx = os.path.join(temp_dir, "converted.docx")
    
    system = platform.system()

    # 方法1: 优先使用 LibreOffice (保留样式更完整)
    libreoffice_paths = []
    if system == "Darwin":
        libreoffice_paths = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "/usr/local/bin/soffice",
        ]
    elif system == "Linux":
        libreoffice_paths = [
            "/usr/bin/soffice",
            "/usr/bin/libreoffice",
        ]
    elif system == "Windows":
        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]

    # 尝试从 PATH 中获取 soffice/libreoffice
    for candidate in ("soffice", "libreoffice"):
        found = shutil.which(candidate)
        if found:
            libreoffice_paths.insert(0, found)
    
    for soffice_path in libreoffice_paths:
        if os.path.exists(soffice_path):
            try:
                result = subprocess.run(
                    [
                        soffice_path,
                        "--headless",
                        "--convert-to", "docx",
                        "--outdir", temp_dir,
                        doc_path
                    ],
                    capture_output=True,
                    text=True,
                    timeout=120
                )
                # LibreOffice 输出文件名基于原文件名
                base_name = os.path.splitext(os.path.basename(doc_path))[0]
                converted_file = os.path.join(temp_dir, f"{base_name}.docx")
                if os.path.exists(converted_file):
                    # 重命名为统一的临时文件名
                    shutil.move(converted_file, temp_docx)
                    print(f"📎 使用 LibreOffice 将 .doc 转换为 .docx")
                    return temp_docx
            except (subprocess.TimeoutExpired, FileNotFoundError):
                continue

    # 方法2: macOS 使用 textutil（样式保留较弱，作为兜底）
    if system == "Darwin":
        try:
            result = subprocess.run(
                ["textutil", "-convert", "docx", doc_path, "-output", temp_docx],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0 and os.path.exists(temp_docx):
                print(f"📎 使用 textutil 将 .doc 转换为 .docx")
                return temp_docx
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
    
    # 清理临时目录
    shutil.rmtree(temp_dir)
    
    raise RuntimeError(
        "无法转换 .doc 文件：未找到可用的转换工具。\n"
        "请安装以下任一工具：\n"
        "  - macOS: 系统自带 textutil（应该已可用）\n"
        "  - 跨平台: LibreOffice (https://www.libreoffice.org/)\n"
        "或者使用 Microsoft Word 将文件另存为 .docx 格式后再转换。"
    )


def convert_docm_to_docx(docm_path: str) -> str:
    """
    将DOCM文件转换为DOCX格式（临时文件）
    DOCM是启用了宏的Word文档，需要移除宏才能被python-docx处理
    
    Args:
        docm_path: DOCM文件路径
        
    Returns:
        临时DOCX文件路径
    """
    # 创建临时目录
    temp_dir = tempfile.mkdtemp()
    temp_docx = os.path.join(temp_dir, "converted.docx")
    
    # 复制文件内容，但不包含宏
    with zipfile.ZipFile(docm_path, 'r') as zip_in:
        with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for item in zip_in.infolist():
                data = zip_in.read(item.filename)
                # 跳过宏相关文件
                if item.filename.startswith('word/vbaProject.bin'):
                    continue
                if item.filename == '[Content_Types].xml':
                    # 修改Content_Types.xml，将宏文档类型改为普通文档类型
                    content = data.decode('utf-8')
                    content = content.replace(
                        'application/vnd.ms-word.document.macroEnabled.main+xml',
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'
                    )
                    data = content.encode('utf-8')
                zip_out.writestr(item, data)
    
    return temp_docx


class DocxToMarkdownConverter:
    """DOCX/DOC转Markdown转换器"""

    XML_NAMESPACES = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "v": "urn:schemas-microsoft-com:vml",
    }

    IMAGE_EXT_BY_CONTENT_TYPE = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/tiff": ".tif",
        "image/x-emf": ".emf",
        "image/x-wmf": ".wmf",
        "image/svg+xml": ".svg",
    }

    def __init__(self, docx_path: str, use_heuristic_heading: bool = False):
        """
        初始化转换器
        
        Args:
            docx_path: 文档文件路径（支持 .docx、.docm、.doc 格式）
            use_heuristic_heading: 是否启用基于字体大小/加粗的标题推断
        """
        self.original_path = docx_path
        self.temp_dir = None
        self.temp_docx = None
        self.use_heuristic_heading = use_heuristic_heading
        self.output_dir = None
        self.output_basename = None
        self.image_dir_name = None
        self.image_output_dir = None
        self.image_rel_dir = None
        self.image_map = {}
        self.image_counter = 1
        
        # 获取文件扩展名
        _, ext = os.path.splitext(docx_path.lower())
        
        # 处理旧版 .doc 格式
        if ext == '.doc':
            print(f"🔄 检测到旧版 .doc 格式，正在转换...")
            self.temp_docx = convert_doc_to_docx(docx_path)
            self.temp_dir = os.path.dirname(self.temp_docx)
            self.doc = Document(self.temp_docx)
            self.markdown_lines = []
            return
        
        # 处理 .docx 和 .docm 格式
        try:
            self.doc = Document(docx_path)
        except ValueError as e:
            if 'macroEnabled' in str(e):
                # 尝试转换为普通DOCX（处理宏启用文档）
                self.temp_docx = convert_docm_to_docx(docx_path)
                self.temp_dir = os.path.dirname(self.temp_docx)
                self.doc = Document(self.temp_docx)
            else:
                raise
        
        self.markdown_lines = []
    
    def cleanup(self):
        """清理临时文件"""
        if self.temp_dir and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def _set_output_context(self, output_path: str):
        """设置输出上下文（用于图片导出）"""
        output_dir = os.path.dirname(output_path) or "."
        output_basename = os.path.splitext(os.path.basename(output_path))[0]
        image_dir_name = f"{output_basename}_images"

        self.output_dir = output_dir
        self.output_basename = output_basename
        self.image_dir_name = image_dir_name
        self.image_output_dir = os.path.join(output_dir, image_dir_name)
        self.image_rel_dir = os.path.relpath(self.image_output_dir, output_dir)
        self.image_map = {}
        self.image_counter = 1

    def _ensure_image_output_dir(self) -> bool:
        """确保图片输出目录存在（延迟创建）"""
        if not self.image_output_dir:
            return False
        if not os.path.exists(self.image_output_dir):
            os.makedirs(self.image_output_dir, exist_ok=True)
        return True

    def _get_image_extension(self, part) -> str:
        """获取图片文件扩展名"""
        partname = str(getattr(part, "partname", ""))
        ext = os.path.splitext(partname)[1].lower()
        if ext:
            return ext
        content_type = getattr(part, "content_type", "")
        return self.IMAGE_EXT_BY_CONTENT_TYPE.get(content_type, ".img")

    def _save_image_from_rid(self, r_id: str) -> str:
        """根据关系ID导出图片，返回相对路径"""
        if not r_id:
            return "", ""
        rel = self.doc.part.rels.get(r_id)
        if rel is None or rel.is_external:
            return "", ""
        part = rel.target_part
        if not hasattr(part, "blob"):
            return "", ""
        part_key = str(getattr(part, "partname", r_id))
        if part_key in self.image_map:
            image_info = self.image_map[part_key]
            return image_info.get("path", ""), image_info.get("label", "")

        if not self._ensure_image_output_dir():
            return "", ""

        ext = self._get_image_extension(part)
        index = self.image_counter
        filename = f"image_{index:03d}{ext}"
        label = f"图{index:03d}"
        self.image_counter += 1

        abs_path = os.path.join(self.image_output_dir, filename)
        with open(abs_path, "wb") as f:
            f.write(part.blob)

        rel_path = os.path.join(self.image_rel_dir, filename).replace(os.sep, "/")
        self.image_map[part_key] = {"path": rel_path, "label": label}
        return rel_path, label

    def _get_image_rids_from_run(self, run) -> list:
        """从run中提取图片关系ID"""
        r_ids = []
        blip_tag = qn("a:blip")
        vml_tag = "{urn:schemas-microsoft-com:vml}imagedata"
        for node in run._element.iter():
            if node.tag == blip_tag:
                r_id = node.get(qn("r:embed"))
                if r_id:
                    r_ids.append(r_id)
            elif node.tag == vml_tag:
                r_id = node.get(qn("r:id"))
                if r_id:
                    r_ids.append(r_id)
        return r_ids

    def _collect_run_tokens(self, paragraph: Paragraph) -> list:
        """收集段落中的文本与图片，保持顺序"""
        tokens = []
        for run in paragraph.runs:
            if run.text:
                tokens.append(("text", run.text))
            for r_id in self._get_image_rids_from_run(run):
                rel_path, label = self._save_image_from_rid(r_id)
                if rel_path:
                    tokens.append(("image", rel_path, label))
        return tokens

    def _tokens_to_text(self, tokens: list) -> str:
        """将文本/图片tokens转换为Markdown行内文本"""
        if not tokens:
            return ""
        parts = []
        for idx, token in enumerate(tokens):
            kind = token[0]
            if kind == "text":
                parts.append(token[1])
                continue
            if parts and not parts[-1].endswith((" ", "\n", "\t")):
                parts.append(" ")
            path = token[1]
            label = token[2] if len(token) > 2 else "image"
            parts.append(f"![{label}]({path})")
            if idx + 1 < len(tokens):
                next_token = tokens[idx + 1]
                if next_token[0] == "text" and not next_token[1].startswith((" ", "\n", "\t")):
                    parts.append(" ")
        return "".join(parts)

    def _image_markdown(self, path: str, label: str) -> str:
        """生成带编号的图片Markdown"""
        return f"![{label}]({path})"
        
    def _extract_heading_level_from_text(self, style_text: str) -> int:
        """从样式名称或ID中提取标题级别"""
        if not style_text:
            return 0
        text = style_text.lower()
        match = re.search(r"(heading|标题)\s*([1-6])", text)
        if match:
            return int(match.group(2))
        match = re.search(r"(heading|标题)([1-6])", text)
        if match:
            return int(match.group(2))
        return 0

    def _get_outline_level_from_ppr(self, ppr) -> int:
        """从段落属性中获取 outline level (0-based)"""
        try:
            outline_lvl = getattr(ppr, "outlineLvl", None)
            if outline_lvl is None:
                return 0
            val = getattr(outline_lvl, "val", None)
            if val is None:
                return 0
            level = int(val)
            if 0 <= level <= 5:
                return level + 1
        except (ValueError, TypeError):
            return 0
        return 0

    def _get_heading_level_from_style(self, paragraph: Paragraph) -> int:
        """优先从样式名称/ID中提取标题级别"""
        style = paragraph.style
        if not style:
            return 0
        level = self._extract_heading_level_from_text(style.name)
        if level:
            return level
        level = self._extract_heading_level_from_text(getattr(style, "style_id", ""))
        if level:
            return level
        # 兼容部分样式仅设置了 outline level
        try:
            style_ppr = getattr(style._element, "pPr", None)
            if style_ppr is not None:
                level = self._get_outline_level_from_ppr(style_ppr)
                if level:
                    return level
        except Exception:
            return 0
        return 0

    def _get_heading_level_from_paragraph(self, paragraph: Paragraph) -> int:
        """从段落属性或样式中获取标题级别"""
        # 1) 先看段落自身 outline level
        try:
            ppr = paragraph._p.pPr
            if ppr is not None:
                level = self._get_outline_level_from_ppr(ppr)
                if level:
                    return level
        except Exception:
            pass
        # 2) 再看样式
        return self._get_heading_level_from_style(paragraph)

    def get_paragraph_style_level(self, paragraph: Paragraph) -> tuple:
        """
        获取段落的样式级别
        
        Args:
            paragraph: 段落对象
            
        Returns:
            (is_heading, level): 是否为标题，标题级别(1-6)
        """
        level = self._get_heading_level_from_paragraph(paragraph)
        if level:
            return (True, level)

        if not self.use_heuristic_heading:
            return (False, 0)

        # 基于格式判断标题级别（字体大小和加粗）
        if paragraph.runs:
            first_run = paragraph.runs[0]
            is_bold = first_run.font.bold
            font_size = first_run.font.size

            # 如果段落加粗且字体较大，判断为标题
            if is_bold and font_size:
                # 字体大小单位是 twips (1/20 point)
                # 转换为磅值进行判断
                size_pt = font_size.pt if hasattr(font_size, "pt") else font_size / 12700

                # 根据字体大小判断标题级别
                if size_pt >= 18:  # 一级标题
                    return (True, 1)
                elif size_pt >= 16:  # 二级标题
                    return (True, 2)
                elif size_pt >= 14:  # 三级标题
                    return (True, 3)
                elif size_pt >= 12:  # 四级标题
                    return (True, 4)
                elif size_pt >= 10:  # 五级标题
                    return (True, 5)
                elif size_pt >= 8:  # 六级标题
                    return (True, 6)

        return (False, 0)
    
    def convert_paragraph(self, paragraph: Paragraph) -> str:
        """
        转换段落为Markdown格式
        
        Args:
            paragraph: 段落对象
            
        Returns:
            Markdown格式的文本
        """
        tokens = self._collect_run_tokens(paragraph)
        if not tokens:
            return ""

        is_heading, level = self.get_paragraph_style_level(paragraph)
        
        if is_heading:
            # 标题格式
            heading_text = paragraph.text.strip()
            if heading_text:
                md = f"{'#' * level} {heading_text}\n"
                image_tokens = [token for token in tokens if token[0] == "image"]
                if image_tokens:
                    md += "\n".join(
                        [self._image_markdown(token[1], token[2]) for token in image_tokens]
                    ) + "\n"
                return md

        # 普通正文
        content = self._tokens_to_text(tokens).strip()
        if not content:
            return ""
        return f"{content}\n"
    
    def convert_table(self, table: Table) -> str:
        """
        转换表格为Markdown格式
        
        Args:
            table: 表格对象
            
        Returns:
            Markdown格式的表格
        """
        if not table.rows:
            return ""
        
        markdown_table = []
        
        # 处理表格的每一行
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for cell in row.cells:
                # 获取单元格文本，处理多行内容（含图片）
                cell_parts = []
                for paragraph in cell.paragraphs:
                    tokens = self._collect_run_tokens(paragraph)
                    if not tokens:
                        continue
                    cell_text = self._tokens_to_text(tokens).strip()
                    if cell_text:
                        cell_parts.append(cell_text)
                cell_text = "<br>".join(cell_parts)
                row_cells.append(cell_text)
            
            # 添加表格行
            markdown_table.append("| " + " | ".join(row_cells) + " |")
            
            # 在第一行后添加分隔符
            if row_idx == 0:
                separator = "| " + " | ".join(["---"] * len(row_cells)) + " |"
                markdown_table.append(separator)
        
        return "\n".join(markdown_table) + "\n"
    
    def convert(self) -> str:
        """
        执行转换
        
        Returns:
            Markdown格式的文本
        """
        self.markdown_lines = []
        
        # 遍历文档中的所有元素
        for element in self.doc.element.body:
            if isinstance(element, CT_P):
                # 段落元素
                paragraph = Paragraph(element, self.doc)
                md_text = self.convert_paragraph(paragraph)
                if md_text:
                    self.markdown_lines.append(md_text)
            
            elif isinstance(element, CT_Tbl):
                # 表格元素
                table = Table(element, self.doc)
                md_table = self.convert_table(table)
                if md_table:
                    self.markdown_lines.append(md_table)
                    self.markdown_lines.append("")  # 表格后添加空行
        
        return "\n".join(self.markdown_lines)
    
    def save_to_file(self, output_path: str):
        """
        保存转换结果到文件
        
        Args:
            output_path: 输出文件路径
        """
        self._set_output_context(output_path)
        markdown_content = self.convert()
        
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        print(f"✅ 转换完成！")
        print(f"📄 输入文件: {self.original_path}")
        print(f"📝 输出文件: {output_path}")
        print(f"📊 文档统计:")
        print(f"   - 段落数: {len(self.doc.paragraphs)}")
        print(f"   - 表格数: {len(self.doc.tables)}")
        print(f"   - 图片数: {len(self.image_map)}")


def main():
    """主函数"""
    default_input_file = "/Users/zhangxy/1/1.docx"

    parser = argparse.ArgumentParser(
        description="将Word文档转换为Markdown格式，保留表格和标题层级关系。支持 .docx、.docm、.doc 格式。"
    )
    parser.add_argument(
        "input_file",
        nargs="?",
        default=default_input_file,
        help="输入Word文件路径，支持 .docx/.docm/.doc 格式（默认使用脚本内置示例文件）",
    )
    parser.add_argument(
        "-o",
        "--output",
        dest="output_file",
        default=None,
        help="输出 Markdown 文件路径（默认与输入文件同目录同名，扩展名为 .md）",
    )
    parser.add_argument(
        "--heuristic-heading",
        action="store_true",
        help="启用基于字体大小/加粗的标题推断（可能导致标题层级与原文不一致）",
    )
    args = parser.parse_args()

    input_file = args.input_file
    output_file = args.output_file or (os.path.splitext(input_file)[0] + ".md")
    
    # 检查输入文件是否存在
    if not os.path.exists(input_file):
        print(f"❌ 错误: 文件不存在 - {input_file}")
        return
    
    converter = None
    try:
        # 创建转换器并执行转换
        converter = DocxToMarkdownConverter(
            input_file,
            use_heuristic_heading=args.heuristic_heading
        )
        converter.save_to_file(output_file)
        
    except Exception as e:
        print(f"❌ 转换失败: {str(e)}")
        import traceback
        traceback.print_exc()
    finally:
        # 清理临时文件
        if converter:
            converter.cleanup()


if __name__ == "__main__":
    main()
