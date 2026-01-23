"""将 Markdown 文件转换为 DOCX 文档，保留标题层级和格式。"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


def parse_markdown_line(line: str) -> Tuple[str, str]:
    """
    解析 Markdown 行，返回 (类型, 内容)。
    
    类型可能是: 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'list', 'text'
    """
    line = line.rstrip()
    
    # 匹配标题 (# 到 ######)
    heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
    if heading_match:
        level = len(heading_match.group(1))
        content = heading_match.group(2)
        return f'h{level}', content
    
    # 匹配列表项 (- 或 * 开头)
    list_match = re.match(r'^[\-\*]\s+(.+)$', line)
    if list_match:
        content = list_match.group(1)
        return 'list', content
    
    # 普通文本
    if line.strip():
        return 'text', line
    
    # 空行
    return 'empty', ''


def set_chinese_font(run, font_name: str = '宋体'):
    """设置中文字体（解决 python-docx 中文字体问题）。"""
    run.font.name = font_name
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def create_numbering(doc: Document):
    """创建多级自动编号并链接到标题样式."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn

    # 直接修改 styles.xml 中的 numbering 定义
    styles = doc.styles

    # 为每个标题级别设置正确的 numPr
    for i in range(1, 10):
        style_name = f'Heading {i}'
        if style_name in styles:
            style = styles[style_name]
            # 获取或创建 pPr
            if style._element.pPr is None:
                style._element.get_or_add_pPr()

            # 移除旧的 numPr
            old_numPr = style._element.pPr.find(qn('w:numPr'))
            if old_numPr is not None:
                style._element.pPr.remove(old_numPr)

            # 创建新的 numPr
            numPr = OxmlElement('w:numPr')
            style._element.pPr.append(numPr)

            # 创建 ilvl 元素
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), str(i - 1))
            numPr.append(ilvl)

            # 创建 numId 元素
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), str(i + 1))  # 使用与标题级别相同的 numId（避开 numId=1 用于bullet）
            numPr.append(numId)

    # 更新 List Bullet 样式的 numId
    if 'List Bullet' in styles:
        list_style = styles['List Bullet']
        if list_style._element.pPr is not None:
            old_numPr = list_style._element.pPr.find(qn('w:numPr'))
            if old_numPr is not None:
                list_style._element.pPr.remove(old_numPr)
        if list_style._element.pPr is None:
            list_style._element.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        list_style._element.pPr.append(numPr)
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')  # 使用 numId=1（bullet 编号）
        numPr.append(numId)


def fix_numbering_after_save(docx_path: Path):
    """保存后修复 numbering.xml."""
    import zipfile
    import shutil

    # 新的 multilevel numbering XML
    new_numbering = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
    <w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:abstractNum w:abstractNumId="0">
            <w:multiLevelType w:val="multilevel"/>
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="0" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:eastAsia="宋体"/>
                    <w:b/>
                    <w:sz w:val="44"/>
                    <w:color w:val="000000"/>
                </w:rPr>
                <w:suff w:val="space"/>
            </w:lvl>
            <w:lvl w:ilvl="1">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1.%2."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="0" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:eastAsia="宋体"/>
                    <w:b/>
                    <w:sz w:val="36"/>
                    <w:color w:val="000000"/>
                </w:rPr>
                <w:suff w:val="space"/>
            </w:lvl>
            <w:lvl w:ilvl="2">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1.%2.%3."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="0" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:eastAsia="宋体"/>
                    <w:b/>
                    <w:sz w:val="32"/>
                    <w:color w:val="000000"/>
                </w:rPr>
                <w:suff w:val="space"/>
            </w:lvl>
            <w:lvl w:ilvl="3">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1.%2.%3.%4."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="0" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:eastAsia="宋体"/>
                    <w:b/>
                    <w:sz w:val="28"/>
                    <w:color w:val="000000"/>
                </w:rPr>
                <w:suff w:val="space"/>
            </w:lvl>
            <w:lvl w:ilvl="4">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1.%2.%3.%4.%5."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="0" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:eastAsia="宋体"/>
                    <w:b/>
                    <w:sz w:val="24"/>
                    <w:color w:val="000000"/>
                </w:rPr>
                <w:suff w:val="space"/>
            </w:lvl>
            <w:lvl w:ilvl="5">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1.%2.%3.%4.%5.%6."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="0" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:eastAsia="宋体"/>
                    <w:b/>
                    <w:sz w:val="24"/>
                    <w:color w:val="000000"/>
                </w:rPr>
                <w:suff w:val="space"/>
            </w:lvl>
            <w:lvl w:ilvl="6">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1.%2.%3.%4.%5.%6.%7."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="0" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:eastAsia="宋体"/>
                    <w:b/>
                    <w:sz w:val="24"/>
                    <w:color w:val="000000"/>
                </w:rPr>
                <w:suff w:val="space"/>
            </w:lvl>
            <w:lvl w:ilvl="7">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1.%2.%3.%4.%5.%6.%7.%8."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="0" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:eastAsia="宋体"/>
                    <w:b/>
                    <w:sz w:val="24"/>
                    <w:color w:val="000000"/>
                </w:rPr>
                <w:suff w:val="space"/>
            </w:lvl>
            <w:lvl w:ilvl="8">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1.%2.%3.%4.%5.%6.%7.%8.%9."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="0" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:eastAsia="宋体"/>
                    <w:b/>
                    <w:sz w:val="24"/>
                    <w:color w:val="000000"/>
                </w:rPr>
                <w:suff w:val="space"/>
            </w:lvl>
        </w:abstractNum>
        <!-- Bullet 编号定义 -->
        <w:abstractNum w:abstractNumId="1">
            <w:multiLevelType w:val="multilevel"/>
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="bullet"/>
                <w:lvlText w:val="&#x2022;"/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="480" w:hanging="0"/>
                </w:pPr>
                <w:rPr>
                    <w:rFonts w:eastAsia="宋体" w:ascii="Symbol" w:hAnsi="Symbol"/>
                    <w:sz w:val="24"/>
                    <w:color w:val="000000"/>
                </w:rPr>
                <w:suff w:val="space"/>
            </w:lvl>
        </w:abstractNum>
        <w:num w:numId="1">
            <w:abstractNumId w:val="1"/>
        </w:num>
        <w:num w:numId="2">
            <w:abstractNumId w:val="0"/>
        </w:num>
        <w:num w:numId="3">
            <w:abstractNumId w:val="0"/>
        </w:num>
        <w:num w:numId="4">
            <w:abstractNumId w:val="0"/>
        </w:num>
        <w:num w:numId="5">
            <w:abstractNumId w:val="0"/>
        </w:num>
        <w:num w:numId="6">
            <w:abstractNumId w:val="0"/>
        </w:num>
        <w:num w:numId="7">
            <w:abstractNumId w:val="0"/>
        </w:num>
        <w:num w:numId="8">
            <w:abstractNumId w:val="0"/>
        </w:num>
        <w:num w:numId="9">
            <w:abstractNumId w:val="0"/>
        </w:num>
        <w:num w:numId="10">
            <w:abstractNumId w:val="0"/>
        </w:num>
    </w:numbering>'''

    # 创建临时文件
    temp_path = docx_path.with_suffix('.temp.docx')

    # 读取原文件，替换 numbering.xml
    with zipfile.ZipFile(docx_path, 'r') as zf_in:
        with zipfile.ZipFile(temp_path, 'w', compression=zipfile.ZIP_DEFLATED) as zf_out:
            for item in zf_in.infolist():
                if item.filename != 'word/numbering.xml':
                    # 复制其他文件
                    data = zf_in.read(item.filename)
                    zf_out.writestr(item, data)
            # 写入新的 numbering.xml
            zf_out.writestr('word/numbering.xml', new_numbering)

    # 替换原文件
    shutil.move(temp_path, docx_path)


def add_paragraph_with_style(
    doc: Document,
    content: str,
    style_type: str,
    font_name: str = '宋体',
    font_size: int = 12,
    bold: bool = False,
    color: RGBColor = RGBColor(0, 0, 0)
):
    """添加带样式的段落。"""
    
    # 根据类型选择样式（使用预设标题样式以保留大纲级别）
    style_map = {
        'h1': 'Heading 1',
        'h2': 'Heading 2',
        'h3': 'Heading 3',
        'h4': 'Heading 4',
        'h5': 'Heading 5',
        'h6': 'Heading 6',
        'list': 'List Bullet',
        'text': 'Normal',
    }

    style = style_map.get(style_type, 'Normal')
    para = doc.add_paragraph(content, style=style)

    # 为列表段落手动添加 numPr
    if style_type == 'list':
        # 确保 pPr 存在
        if para._element.pPr is None:
            para._element.get_or_add_pPr()
        # 移除旧的 numPr
        old_numPr = para._element.pPr.find(qn('w:numPr'))
        if old_numPr is not None:
            para._element.pPr.remove(old_numPr)
        # 创建新的 numPr
        numPr = OxmlElement('w:numPr')
        para._element.pPr.append(numPr)
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')  # bullet 编号
        numPr.append(numId)
    
    # 设置字体样式
    for run in para.runs:
        run.font.size = Pt(font_size)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = False  # 覆盖Heading 4自带斜体
        set_chinese_font(run, font_name)

    # 所有标题段后间距为0，正文段前段后间距为0
    para_format = para.paragraph_format
    if style_type in ('h1', 'h2', 'h3'):
        para_format.space_before = Pt(6)  # 0.5行 ≈ 6磅
        para_format.space_after = Pt(0)
    elif style_type.startswith('h'):
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(0)
    elif style_type == 'text':
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(0)

    return para


def convert_markdown_to_docx(
    markdown_path: Path,
    output_path: Path | None = None,
    font_name: str = '宋体',
    font_color: RGBColor = RGBColor(0, 0, 0)
) -> Path:
    """
    将 Markdown 文件转换为 DOCX 文档。
    
    Args:
        markdown_path: Markdown 文件路径
        output_path: 输出 DOCX 文件路径（默认与 Markdown 同名）
        font_name: 字体名称（默认宋体）
        font_color: 字体颜色（默认黑色）
    
    Returns:
        生成的 DOCX 文件路径
    """
    markdown_path = markdown_path.expanduser().resolve()
    if not markdown_path.exists():
        raise FileNotFoundError(f"Markdown 文件不存在: {markdown_path}")
    
    # 确定输出路径
    if output_path is None:
        output_path = markdown_path.with_suffix('.docx')
    else:
        output_path = output_path.expanduser().resolve()
    
    # 创建 Word 文档
    doc = Document()

    # 创建多级编号
    create_numbering(doc)

    # 读取 Markdown 内容
    content = markdown_path.read_text(encoding='utf-8')
    lines = content.split('\n')
    
    # 逐行解析并添加到文档
    for line in lines:
        line_type, line_content = parse_markdown_line(line)
        
        if line_type == 'empty':
            # 空行，跳过不添加段落
            continue
        
        # 根据类型设置字体大小和加粗
        if line_type == 'h1':
            font_size = 22
            bold = True
        elif line_type == 'h2':
            font_size = 18
            bold = True
        elif line_type == 'h3':
            font_size = 16
            bold = True
        elif line_type == 'h4':
            font_size = 14
            bold = True
        elif line_type in ('h5', 'h6'):
            font_size = 12
            bold = True
        else:
            font_size = 12
            bold = False
        
        # 添加段落
        add_paragraph_with_style(
            doc,
            line_content,
            line_type,
            font_name=font_name,
            font_size=font_size,
            bold=bold,
            color=font_color
        )

    # 保存文档
    doc.save(str(output_path))

    # 修复 numbering.xml（添加多级编号）
    fix_numbering_after_save(output_path)

    print(f"✅ 已生成 DOCX 文档: {output_path}")
    
    return output_path


def main():
    """命令行入口。"""
    parser = argparse.ArgumentParser(
        description="将 Markdown 文件转换为 DOCX 文档，保留标题层级和格式"
    )
    parser.add_argument(
        'markdown_path',
        type=Path,
        help='待转换的 Markdown 文件路径'
    )
    parser.add_argument(
        '-o', '--output',
        type=Path,
        default=None,
        help='输出 DOCX 文件路径（默认与 Markdown 同名）'
    )
    parser.add_argument(
        '-f', '--font',
        type=str,
        default='宋体',
        help='字体名称（默认：宋体）'
    )
    
    args = parser.parse_args()
    
    # 执行转换
    convert_markdown_to_docx(
        markdown_path=args.markdown_path,
        output_path=args.output,
        font_name=args.font
    )


if __name__ == '__main__':
    main()
